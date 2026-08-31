"""ci_extract_events — extract CI events from user-provided documents."""

from __future__ import annotations

import re
from typing import Any

from pydantic import BaseModel, Field


class ExtractEventsArgs(BaseModel):
    file_path: str = Field(description="Path to document: PDF, DOCX, TXT, CSV, or MD")
    focus: list[str] | None = Field(
        default=None,
        description="Focus areas: trial_readout, regulatory, label_change, pricing, market_access, partnership",
    )
    therapeutic_area: str | None = Field(
        default=None,
        description="Therapeutic area context for relevance filtering",
    )
    max_pages: int = Field(default=50, ge=1, le=200, description="Max pages to process (PDF only)")


TOOL: dict[str, Any] = {
    "name": "ci_extract_events",
    "description": (
        "Extract structured competitive intelligence events from a document "
        "(PDF, DOCX, TXT, CSV, Markdown). Uses heuristic pattern matching "
        "to identify dates, company names, drug names, and event types. "
        "Returns events in the common CIEvent schema."
    ),
    "args": ExtractEventsArgs,
}

_PHARMA_COMPANIES = {
    "pfizer",
    "merck",
    "roche",
    "novartis",
    "johnson & johnson",
    "j&j",
    "astrazeneca",
    "bristol-myers squibb",
    "bms",
    "eli lilly",
    "lilly",
    "abbvie",
    "amgen",
    "gilead",
    "sanofi",
    "gsk",
    "glaxosmithkline",
    "bayer",
    "takeda",
    "regeneron",
    "biogen",
    "moderna",
    "vertex",
    "seagen",
    "incyte",
    "jazz",
    "daiichi sankyo",
    "astellas",
    "boehringer ingelheim",
    "merck kgaa",
    "shire",
    "alexion",
}

_EVENT_PATTERNS = {
    "approval": re.compile(r"(?:FDA|EMA|NMPA|PMDA|TGA)\s+(?:approved|grants?\s+approval|clearance)", re.IGNORECASE),
    "trial_readout": re.compile(
        r"(?:phase\s+[I1-4]+|pivotal)\s+(?:trial|study)\s+(?:met|missed|results?|data|readout)", re.IGNORECASE
    ),
    "regulatory_submission": re.compile(
        r"(?:filed|submitted|submission|sNDA|sBLA|MAA|NDA|BLA)\s+(?:to|with|for)?\s*(?:FDA|EMA)?", re.IGNORECASE
    ),
    "label_change": re.compile(
        r"label\s+(?:change|update|revision|expansion)|new\s+indication|supplemental\s+approval", re.IGNORECASE
    ),
    "partnership": re.compile(
        r"(?:partnership|collaboration|licensing|co-development|acquisition|merger|deal)\s+(?:with|between|agreement)?",
        re.IGNORECASE,
    ),
    "pricing": re.compile(
        r"(?:list\s+price|WAC|launch\s+price|pricing\s+(?:strategy|decision|announcement)|reimbursement\s+(?:decision|approval))",
        re.IGNORECASE,
    ),
    "discontinuation": re.compile(
        r"(?:discontinue|terminate|halt|suspend|withdraw)\w*\s+(?:trial|study|program|development)", re.IGNORECASE
    ),
}

_DATE_PATTERN = re.compile(
    r"(?:(?:January|February|March|April|May|June|July|August|September|October|November|December)"
    r"\s+\d{1,2},?\s+\d{4})"
    r"|(?:\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)"
    r"\s+\d{4})"
    r"|(?:\d{4}-\d{2}-\d{2})"
    r"|(?:\d{1,2}/\d{1,2}/\d{4})",
    re.IGNORECASE,
)


def handle(arguments: dict[str, Any]) -> list[dict[str, Any]]:
    import json
    import os

    file_path = arguments["file_path"]
    focus = arguments.get("focus")
    therapeutic_area = arguments.get("therapeutic_area")
    max_pages = arguments.get("max_pages", 50)

    if not os.path.isfile(file_path):
        return [{"type": "text", "text": json.dumps({"error": f"File not found: {file_path}"})}]

    ext = os.path.splitext(file_path)[1].lower()
    pages = _extract_text(file_path, ext, max_pages)

    events = []
    for page in pages:
        page_events = _extract_events_from_text(page["text"], page["page_number"], os.path.basename(file_path), focus)
        events.extend(page_events)

    if therapeutic_area:
        for e in events:
            e["therapeutic_area"] = therapeutic_area

    events.sort(key=lambda e: e.get("date") or "", reverse=True)

    output = {
        "file": os.path.basename(file_path),
        "pages_processed": len(pages),
        "total_events": len(events),
        "therapeutic_area": therapeutic_area,
        "events": events,
    }
    return [{"type": "text", "text": json.dumps(output, indent=2)}]


def _extract_text(file_path: str, ext: str, max_pages: int) -> list[dict[str, Any]]:
    if ext == ".pdf":
        return _extract_pdf(file_path, max_pages)
    if ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    return _extract_plain(file_path)


def _extract_pdf(file_path: str, max_pages: int) -> list[dict[str, Any]]:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return [{"page_number": 1, "text": _read_file_text(file_path)}]

    doc = pdfium.PdfDocument(file_path)
    pages = []
    for i in range(min(len(doc), max_pages)):
        textpage = doc[i].get_textpage()
        text = textpage.get_text_bounded()
        textpage.close()
        if text.strip():
            pages.append({"page_number": i + 1, "text": text})
    doc.close()
    return pages


def _extract_docx(file_path: str) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError:
        return [{"page_number": 1, "text": _read_file_text(file_path)}]

    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [{"page_number": 1, "text": text}]


def _extract_plain(file_path: str) -> list[dict[str, Any]]:
    return [{"page_number": 1, "text": _read_file_text(file_path)}]


def _read_file_text(file_path: str) -> str:
    try:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception:
        return ""


def _extract_events_from_text(
    text: str, page_number: int, source: str, focus: list[str] | None
) -> list[dict[str, Any]]:
    if not text or len(text) < 20:
        return []

    sentences = re.split(r"[.!?\n]+", text)
    events = []

    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 20:
            continue

        event_type = None
        for etype, pattern in _EVENT_PATTERNS.items():
            if pattern.search(sentence):
                event_type = etype
                break

        if event_type is None:
            continue

        if focus and event_type not in focus:
            continue

        competitor = _find_competitor(sentence)
        date = _find_date(sentence)

        events.append(
            {
                "event_type": event_type,
                "date": date,
                "competitor": competitor or "Unknown",
                "product": None,
                "description": sentence[:300],
                "source": source,
                "source_page": page_number,
                "confidence": "medium" if competitor else "low",
            }
        )

    return events


def _find_competitor(text: str) -> str | None:
    text_lower = text.lower()
    for company in _PHARMA_COMPANIES:
        if company in text_lower:
            return company.title()
    return None


def _find_date(text: str) -> str | None:
    m = _DATE_PATTERN.search(text)
    if m:
        return m.group(0)
    return None
